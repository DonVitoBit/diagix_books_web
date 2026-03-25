"""
Экспорт статьи в DOCX.
"""

import io
import os
import re
from typing import Optional, Tuple, List

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


_IMG_RE = re.compile(r"!\[(?P<alt>[^\]]*)\]\((?P<src>[^)]+)\)")
_TABLE_SEP_RE = re.compile(r"^\s*\|?\s*:?-{2,}:?\s*(\|\s*:?-{2,}:?\s*)+\|?\s*$")


def _clean_inline_md(text: str) -> str:
    """Убирает базовые markdown-маркеры внутри строки."""
    if not text:
        return ""
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = text.replace("**", "").replace("__", "")
    text = text.replace("*", "").replace("_", "")
    return text.strip()


def _split_table_row(line: str) -> List[str]:
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [_clean_inline_md(c.strip()) for c in s.split("|")]


def text_to_docx(text: str) -> Tuple[Optional[bytes], Optional[str]]:
    """Конвертирует markdown-like текст статьи в DOCX."""
    if not text or not text.strip():
        return None, "Нет текста для экспорта"

    try:
        doc = Document()

        # Базовый стиль абзаца
        style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)

        lines = text.splitlines()
        i = 0
        while i < len(lines):
            raw = lines[i]
            line = raw.strip()
            if not line:
                doc.add_paragraph("")
                i += 1
                continue

            if line == "---":
                i += 1
                continue

            # Markdown table:
            # | col1 | col2 |
            # | ---  | ---  |
            # | ...  | ...  |
            if "|" in line and i + 1 < len(lines) and _TABLE_SEP_RE.match(lines[i + 1].strip() or ""):
                table_rows: List[List[str]] = []
                # header
                table_rows.append(_split_table_row(line))
                # data rows (skip separator row)
                j = i + 2
                while j < len(lines):
                    row_line = lines[j].strip()
                    if not row_line or "|" not in row_line:
                        break
                    table_rows.append(_split_table_row(row_line))
                    j += 1

                max_cols = max(len(r) for r in table_rows) if table_rows else 0
                if max_cols > 0:
                    table = doc.add_table(rows=len(table_rows), cols=max_cols)
                    table.style = "Table Grid"
                    for r_idx, row in enumerate(table_rows):
                        for c_idx in range(max_cols):
                            value = row[c_idx] if c_idx < len(row) else ""
                            table.cell(r_idx, c_idx).text = value
                i = j
                continue

            # Markdown image: ![caption](file:/path.png)
            img_match = _IMG_RE.match(line)
            if img_match:
                alt = _clean_inline_md(img_match.group("alt"))
                src = (img_match.group("src") or "").strip()
                if src.startswith("file:"):
                    src = src[5:]
                if src and os.path.isfile(src):
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(src, width=Inches(5.8))
                    if alt:
                        cap = doc.add_paragraph(alt)
                        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                i += 1
                continue

            # Заголовки
            if line.startswith("# "):
                doc.add_heading(_clean_inline_md(line[2:]), level=1)
                i += 1
                continue
            if line.startswith("## "):
                doc.add_heading(_clean_inline_md(line[3:]), level=2)
                i += 1
                continue
            if line.startswith("### "):
                doc.add_heading(_clean_inline_md(line[4:]), level=3)
                i += 1
                continue

            # Маркированные/нумерованные списки
            if re.match(r"^[-*]\s+", line):
                doc.add_paragraph(_clean_inline_md(re.sub(r"^[-*]\s+", "", line)), style="List Bullet")
                i += 1
                continue
            if re.match(r"^\d+[\.\)]\s+", line):
                doc.add_paragraph(_clean_inline_md(re.sub(r"^\d+[\.\)]\s+", "", line)), style="List Number")
                i += 1
                continue

            # Обычный абзац
            p = doc.add_paragraph(_clean_inline_md(line))
            p.paragraph_format.first_line_indent = Inches(0.25)
            p.paragraph_format.line_spacing = 1.25
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            i += 1

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue(), None
    except Exception as e:
        return None, f"Ошибка DOCX: {str(e)[:200]}"

